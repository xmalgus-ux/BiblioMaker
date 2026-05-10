"""
Модуль для работы с документами Word (.docx)
Извлечение и сохранение списка литературы с сохранением форматирования
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from typing import List, Dict, Tuple


class DocxProcessor:
    """Класс для обработки Word-документов"""

    # Ключевые слова для поиска списка литературы
    BIBLIO_KEYWORDS = [
        r'(?i)список\s+литературы',
        r'(?i)библиографический\s+список',
        r'(?i)использованные\s+источники',
        r'(?i)литература',
        r'(?i)references'
    ]

    def __init__(self, file_path: str):
        """
        Инициализация процессора документов

        Args:
            file_path: Путь к .docx файлу
        """
        self.file_path = file_path
        self.document = Document(file_path)
        self.full_text = self._extract_full_text()

    def _extract_full_text(self) -> str:
        """Извлечение полного текста из документа"""
        return '\n'.join([para.text for para in self.document.paragraphs])

    def find_bibliography_section(self) -> Tuple[int, int]:
        """
        Поиск раздела со списком литературы

        Returns:
            Кортеж (start_index, end_index) параграфов списка литературы
            или (-1, -1) если не найдено
        """
        start_idx = -1
        end_idx = -1

        # Поиск начала раздела
        for i, para in enumerate(self.document.paragraphs):
            if any(re.search(pattern, para.text) for pattern in self.BIBLIO_KEYWORDS):
                start_idx = i + 1
                break

        if start_idx == -1:
            return (-1, -1)

        # Поиск конца раздела (следующий заголовок или пустая строка)
        for i in range(start_idx, len(self.document.paragraphs)):
            para = self.document.paragraphs[i]

            # Пустая строка после списка обычно отделяет следующий раздел.
            # Если дальше снова идет нумерованный/маркированный источник, считаем список продолжающимся.
            if not para.text.strip() and i > start_idx:
                next_text = ""
                for j in range(i + 1, min(i + 5, len(self.document.paragraphs))):
                    candidate = self.document.paragraphs[j].text.strip()
                    if candidate:
                        next_text = candidate
                        break

                if not next_text or not (
                    re.match(r'^\d+[\.\)]\s', next_text) or re.match(r'^[•\-]\s', next_text)
                ):
                    end_idx = i
                    break

        if end_idx == -1:
            end_idx = len(self.document.paragraphs)

        return (start_idx, end_idx)

    def extract_bibliography(self) -> List[str]:
        """
        Извлечение списка литературы

        Returns:
            Список записей литературы
        """
        start_idx, end_idx = self.find_bibliography_section()

        if start_idx == -1:
            return []

        entries = []
        current_entry = ""

        for i in range(start_idx, end_idx):
            para = self.document.paragraphs[i]
            text = para.text.strip()

            if not text:
                if current_entry:
                    entries.append(current_entry)
                    current_entry = ""
                continue

            # Если строка начинается с цифры или маркера - новая запись
            if re.match(r'^\d+[\.\)]\s', text) or re.match(r'^[•\-]\s', text):
                if current_entry:
                    entries.append(current_entry)
                current_entry = text
            else:
                # Продолжение предыдущей записи
                if current_entry:
                    current_entry += " " + text
                else:
                    current_entry = text

        # Добавляем последнюю запись
        if current_entry:
            entries.append(current_entry)

        return entries

    def replace_bibliography(self, new_entries: List[str], output_path: str):
        """
        Замена списка литературы в документе

        Args:
            new_entries: Новый список литературы
            output_path: Путь для сохранения результата
        """
        start_idx, end_idx = self.find_bibliography_section()

        if start_idx == -1:
            # Если раздел не найден - добавляем в конец
            self.document.add_page_break()
            heading = self.document.add_heading("Список литературы", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for entry in new_entries:
                para = self.document.add_paragraph(entry)
                para.paragraph_format.first_line_indent = Cm(1.25)
                para.paragraph_format.line_spacing = Pt(14)
        else:
            anchor = self.document.paragraphs[start_idx - 1] if start_idx > 0 else None
            # Удаляем старые записи
            # Важно: удаляем с конца чтобы индексы не сдвигались
            for i in range(end_idx - 1, start_idx - 1, -1):
                if i < len(self.document.paragraphs):
                    p = self.document.paragraphs[i]._element
                    p.getparent().remove(p)

            # Вставляем новые записи
            for entry in new_entries:
                para = self.document.add_paragraph(entry)
                para.paragraph_format.first_line_indent = Cm(1.25)
                para.paragraph_format.line_spacing = Pt(14)
                if anchor is not None:
                    anchor._p.addnext(para._p)
                    anchor = para

        self.document.save(output_path)

    def get_document_info(self) -> Dict:
        """Получение информации о документе"""
        return {
            'path': self.file_path,
            'name': os.path.basename(self.file_path),
            'paragraphs_count': len(self.document.paragraphs),
            'has_bibliography': self.find_bibliography_section()[0] != -1,
            'bibliography_count': len(self.extract_bibliography())
        }

    @staticmethod
    def create_backup(file_path: str) -> str:
        """
        Создание резервной копии файла

        Args:
            file_path: Путь к исходному файлу

        Returns:
            Путь к резервной копии
        """
        from shutil import copy2
        from datetime import datetime

        base, ext = os.path.splitext(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = f"{base}_backup_{timestamp}{ext}"

        copy2(file_path, backup_path)
        return backup_path
