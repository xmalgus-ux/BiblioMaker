"""
Основное окно приложения BiblioMaker
"""

from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTextEdit, QLabel, QFileDialog,
    QProgressBar, QMessageBox, QSplitter, QGroupBox,
    QComboBox, QTableWidget, QTableWidgetItem, QHeaderView,
    QTabWidget, QLineEdit, QDialog, QDialogButtonBox, QFormLayout,
    QSizePolicy, QSpinBox, QCheckBox,
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QAction, QColor
from PyQt6.QtWidgets import QMenu, QApplication
import os
import re
import sys

# Добавляем корень проекта в path для импортов
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from app.core.database import GOSTDatabase
from app.core.docx_processor import DocxProcessor
from app.core.ai_providers import AIProviderFactory
from app.core.source_classifier import classify_entry, analyze_entry
from typing import List, Optional


class ProcessingThread(QThread):
    """Поток обработки документа"""
    progress = pyqtSignal(int, str)
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)
    cancelled = pyqtSignal(str)

    def __init__(self, file_path: str = None, entries: List[str] = None,
                 prompt_entries: List[str] = None,
                 provider_name: str = None,
                 ui_language: str = "ru_RU"):
        """
        Инициализация потока

        Args:
            file_path: Путь к файлу (для режима файла)
            entries: Список записей (для ручного режима)
            prompt_entries: Записи для промпта (с TYPE=...)
            provider_name: Название ИИ-провайдера
        """
        super().__init__()
        self.file_path = file_path
        self.entries = entries
        self.prompt_entries = prompt_entries
        self.provider_name = provider_name
        self.ui_language = ui_language
        self._cancel_requested = False

    def request_cancel(self):
        """Запрос на отмену обработки"""
        self._cancel_requested = True

    def _check_cancel(self) -> bool:
        if self._cancel_requested:
            self.cancelled.emit("Обработка отменена пользователем")
            return True
        return False

    def run(self):
        try:
            self.progress.emit(10, "Подготовка данных...")
            if self._check_cancel():
                return

            # Получаем записи: из файла или из параметров
            if self.file_path and not self.entries:
                # Режим файла
                processor = DocxProcessor(self.file_path)
                self.progress.emit(30, "Извлечение списка литературы...")
                raw_entries = processor.extract_bibliography()
            else:
                # Режим ручного ввода
                raw_entries = self.entries or []
                processor = None

            if self._check_cancel():
                return

            if not raw_entries:
                self.error.emit("Не найдено записей для обработки")
                return

            self.progress.emit(50, f"Обработка через {self.provider_name}...")

            # Создание провайдера
            provider = AIProviderFactory.create_provider(self.provider_name)

            # Обработка через ИИ
            prompt_entries = self.prompt_entries or raw_entries
            result = provider.fix_text(prompt_entries, ui_language=self.ui_language)

            if self._check_cancel():
                return

            if not result['success']:
                self.error.emit(f"Ошибка ИИ: {result.get('error', 'Неизвестная ошибка')}")
                return

            self.progress.emit(80, "Формирование результата...")

            # Формирование результата
            output = {
                'original': raw_entries,
                'fixed': result['fixed_lines'],
                'provider': result['provider'],
                'processor': processor,
                'sources_count': len(raw_entries),
                'changes_count': sum(1 for i, f in zip(raw_entries, result['fixed_lines']) if i != f),
                'is_manual': processor is None
            }

            self.progress.emit(100, "Готово!")
            self.finished.emit(output)

        except Exception as e:
            import traceback
            self.error.emit(f"Ошибка обработки: {str(e)}\n{traceback.format_exc()}")


class SettingsDialog(QDialog):
    """Диалог настроек приложения"""

    def __init__(self, parent=None, db=None):
        super().__init__(parent)
        self.db = db
        self.initial_settings = {}
        if parent is not None and hasattr(parent, "settings"):
            self.initial_settings.update(parent.settings)
        self.lang = "ru_RU"
        if db:
            self.lang = self.initial_settings.get("ui_language", db.get_setting("ui_language", "ru_RU"))
        self.setMinimumWidth(560)
        self.setup_ui()
        self.load_settings()

    def _texts(self):
        texts = {
            "ru_RU": {
                "title": "Настройки",
                "ui_language": "Язык интерфейса:",
                "theme": "Тема оформления:",
                "default_provider": "ИИ-провайдер:",
                "output_folder": "Папка для сохранения:",
                "auto_backup_row": "Резервная копия:",
                "font_size": "Размер шрифта:",
                "browse": "Выбрать...",
                "choose_folder": "Выберите папку для сохранения",
                "auto_backup": "Создавать резервную копию исходного Word перед сохранением",
                "saved_title": "Успех",
                "saved_text": "Настройки сохранены!",
                "save": "Сохранить",
                "cancel": "Отмена",
                "themes": {"light": "Светлая", "dark": "Темная"},
                "providers": {"mock": "Mock (тестовый)", "yandex": "YandexGPT"},
            },
            "en_US": {
                "title": "Settings",
                "ui_language": "Interface language:",
                "theme": "Theme:",
                "default_provider": "AI provider:",
                "output_folder": "Save folder:",
                "auto_backup_row": "Backup:",
                "font_size": "Font size:",
                "browse": "Browse...",
                "choose_folder": "Select save folder",
                "auto_backup": "Create a backup copy of the original Word file before saving",
                "saved_title": "Success",
                "saved_text": "Settings saved!",
                "save": "Save",
                "cancel": "Cancel",
                "themes": {"light": "Light", "dark": "Dark"},
                "providers": {"mock": "Mock (test)", "yandex": "YandexGPT"},
            },
            "zh_CN": {
                "title": "设置",
                "ui_language": "界面语言:",
                "theme": "主题:",
                "default_provider": "AI 提供商:",
                "output_folder": "保存文件夹:",
                "auto_backup_row": "备份:",
                "font_size": "字体大小:",
                "browse": "选择...",
                "choose_folder": "选择保存文件夹",
                "auto_backup": "保存前为原始 Word 文件创建备份副本",
                "saved_title": "成功",
                "saved_text": "设置已保存!",
                "save": "保存",
                "cancel": "取消",
                "themes": {"light": "浅色", "dark": "深色"},
                "providers": {"mock": "Mock（测试）", "yandex": "YandexGPT"},
            },
        }
        return texts.get(self.lang, texts["ru_RU"])

    def setup_ui(self):
        self.form_layout = QFormLayout(self)

        self.settings_inputs = {}
        self.setting_labels = {}

        language = QComboBox()
        language.addItem("Русский", "ru_RU")
        language.addItem("English", "en_US")
        language.addItem("中文", "zh_CN")
        language.currentIndexChanged.connect(self._on_language_changed)
        self.settings_inputs["ui_language"] = language
        self._add_setting_row("ui_language", language)

        theme = QComboBox()
        self.settings_inputs["theme"] = theme
        self._add_setting_row("theme", theme)

        provider = QComboBox()
        self.settings_inputs["default_provider"] = provider
        self._add_setting_row("default_provider", provider)

        folder_row = QWidget()
        folder_layout = QHBoxLayout(folder_row)
        folder_layout.setContentsMargins(0, 0, 0, 0)
        folder_layout.setSpacing(8)
        output_folder = QLineEdit()
        output_folder.setReadOnly(True)
        self.btn_browse_folder = QPushButton()
        self.btn_browse_folder.clicked.connect(self._choose_output_folder)
        folder_layout.addWidget(output_folder, 1)
        folder_layout.addWidget(self.btn_browse_folder)
        self.settings_inputs["output_folder"] = output_folder
        self._add_setting_row("output_folder", folder_row)

        auto_backup = QCheckBox()
        self.settings_inputs["auto_backup"] = auto_backup
        self._add_setting_row("auto_backup_row", auto_backup)

        font_size = QSpinBox()
        font_size.setRange(9, 18)
        font_size.setSingleStep(1)
        self.settings_inputs["font_size"] = font_size
        self._add_setting_row("font_size", font_size)

        self.buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Save |
            QDialogButtonBox.StandardButton.Cancel
        )
        self.buttons.accepted.connect(self.save_settings)
        self.buttons.rejected.connect(self.reject)
        self.form_layout.addRow(self.buttons)

    def _add_setting_row(self, key: str, widget: QWidget):
        label = QLabel()
        self.setting_labels[key] = label
        self.form_layout.addRow(label, widget)

    def _set_combo_items(self, combo: QComboBox, items: list[tuple[str, str]]):
        current = combo.currentData()
        combo.blockSignals(True)
        combo.clear()
        for text, data in items:
            combo.addItem(text, data)
        index = combo.findData(current)
        if index >= 0:
            combo.setCurrentIndex(index)
        combo.blockSignals(False)

    @staticmethod
    def _fit_button_text(button: QPushButton, extra_padding: int = 46):
        width = button.fontMetrics().horizontalAdvance(button.text()) + extra_padding
        button.setMinimumWidth(max(button.minimumWidth(), width))
        button.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed)

    def _on_language_changed(self):
        self.lang = self.settings_inputs["ui_language"].currentData() or "ru_RU"
        self._apply_language()

    def _apply_language(self):
        text = self._texts()
        self.setWindowTitle(text["title"])
        for key, label in self.setting_labels.items():
            label.setText(text.get(key, text.get(key.replace("_row", ""), key)))
        self.btn_browse_folder.setText(text["browse"])
        self.settings_inputs["auto_backup"].setText(text["auto_backup"])
        self.buttons.button(QDialogButtonBox.StandardButton.Save).setText(text["save"])
        self.buttons.button(QDialogButtonBox.StandardButton.Cancel).setText(text["cancel"])
        self._set_combo_items(
            self.settings_inputs["theme"],
            [(text["themes"]["light"], "light"), (text["themes"]["dark"], "dark")]
        )
        self._set_combo_items(
            self.settings_inputs["default_provider"],
            [(text["providers"]["mock"], "mock"), (text["providers"]["yandex"], "yandex")]
        )
        self._fit_button_text(self.btn_browse_folder)
        self._fit_button_text(self.buttons.button(QDialogButtonBox.StandardButton.Save))
        self._fit_button_text(self.buttons.button(QDialogButtonBox.StandardButton.Cancel))

    def _choose_output_folder(self):
        current = self.settings_inputs["output_folder"].text()
        folder = QFileDialog.getExistingDirectory(
            self,
            self._texts()["choose_folder"],
            current if current and os.path.isdir(current) else ""
        )
        if folder:
            self.settings_inputs["output_folder"].setText(folder)

    def load_settings(self):
        """Загрузка настроек из БД"""
        if not self.db:
            return

        settings = self.db.get_all_settings()
        settings.update({k: v for k, v in self.initial_settings.items() if v is not None})

        language = self.settings_inputs["ui_language"]
        language.blockSignals(True)
        lang_index = language.findData(settings.get("ui_language", self.lang))
        if lang_index >= 0:
            language.setCurrentIndex(lang_index)
        language.blockSignals(False)
        self.lang = language.currentData() or "ru_RU"
        self._apply_language()

        for key, widget in self.settings_inputs.items():
            if key == "ui_language":
                continue
            value = settings.get(key, '')
            if isinstance(widget, QComboBox):
                index = widget.findData(value)
                if index >= 0:
                    widget.setCurrentIndex(index)
            elif isinstance(widget, QCheckBox):
                widget.setChecked(value.lower() == "true")
            elif isinstance(widget, QSpinBox):
                widget.setValue(int(value) if str(value).isdigit() else 12)
            else:
                widget.setText(value)

    def save_settings(self):
        """Сохранение настроек в БД"""
        if not self.db:
            return

        for key, widget in self.settings_inputs.items():
            if isinstance(widget, QComboBox):
                value = widget.currentData()
            elif isinstance(widget, QCheckBox):
                value = "true" if widget.isChecked() else "false"
            elif isinstance(widget, QSpinBox):
                value = str(widget.value())
            else:
                value = widget.text()

            self.db.set_setting(key, value)

        text = self._texts()
        QMessageBox.information(self, text["saved_title"], text["saved_text"])
        self.accept()


class TypeResolutionDialog(QDialog):
    """Диалог выбора типа источников и предупреждений о пропусках"""

    TYPE_OPTIONS = [
        ("unknown", "Неизвестно"),
        ("book", "Книга"),
        ("journal_article", "Статья из журнала"),
        ("collection_article", "Статья из сборника"),
        ("multi_volume", "Многотомное издание"),
        ("gost_standard", "ГОСТ / стандарт"),
        ("federal_law", "Федеральный закон"),
        ("dissertation", "Диссертация / автореферат"),
        ("archive", "Архивный документ"),
        ("electronic", "Электронный ресурс"),
        ("article", "Статья (неуточненная)")
    ]

    def __init__(self, parent=None, issue_rows: List[dict] = None):
        super().__init__(parent)
        self.parent_window = parent
        self.lang = getattr(parent, "settings", {}).get("ui_language", "ru_RU")
        self.setWindowTitle(self._t("title"))
        self.setMinimumWidth(700)
        self.issue_rows = issue_rows or []
        self._type_map = {}
        self._setup_ui()
        self._apply_dialog_theme()

    def _theme_colors(self):
        theme = getattr(self.parent_window, "settings", {}).get("theme", "light")
        if theme == "dark":
            return {
                "input": "#111827",
                "pane": "#1f2937",
                "header": "#374151",
                "text": "#e5e7eb",
                "border": "#4b5563",
                "selection": "#1d4ed8",
                "selection_text": "#ffffff",
            }
        return {
            "input": "#ffffff",
            "pane": "#ffffff",
            "header": "#f1f5f9",
            "text": "#0f172a",
            "border": "#d7dde7",
            "selection": "#bfdbfe",
            "selection_text": "#0f172a",
        }

    def _apply_dialog_theme(self):
        colors = self._theme_colors()
        if self.parent_window is not None:
            self.setStyleSheet(self.parent_window.styleSheet())
        self.table.setStyleSheet("""
            QTableWidget, QTableView {
                background: %(input)s;
                color: %(text)s;
                border: 1px solid %(border)s;
                gridline-color: %(border)s;
                selection-background-color: %(selection)s;
                selection-color: %(selection_text)s;
            }
            QTableWidget::item, QTableView::item {
                background: %(input)s;
                color: %(text)s;
            }
            QTableCornerButton::section,
            QHeaderView::section {
                background: %(header)s;
                color: %(text)s;
                border: 1px solid %(border)s;
            }
            QComboBox {
                background: %(input)s;
                color: %(text)s;
                border: 1px solid %(border)s;
                border-radius: 6px;
                padding: 3px 8px;
            }
            QComboBox QAbstractItemView {
                background: %(pane)s;
                color: %(text)s;
                selection-background-color: %(selection)s;
                selection-color: %(selection_text)s;
            }
        """ % colors)

    def _t(self, key: str):
        texts = {
            "ru_RU": {
                "title": "Уточните тип источников",
                "info": "Некоторые источники не удалось распознать автоматически. Выберите тип для каждой строки.",
                "headers": ["Исходная строка", "Тип источника", "Проблемы"],
                "select_type_title": "Нужно выбрать тип",
                "select_type_text": "Для всех строк выберите тип источника.",
                "types": {
                    "unknown": "Неизвестно", "book": "Книга", "journal_article": "Статья из журнала",
                    "collection_article": "Статья из сборника", "multi_volume": "Многотомное издание",
                    "gost_standard": "ГОСТ / стандарт", "federal_law": "Федеральный закон",
                    "dissertation": "Диссертация / автореферат", "archive": "Архивный документ",
                    "electronic": "Электронный ресурс", "article": "Статья (неуточненная)",
                },
            },
            "en_US": {
                "title": "Clarify Source Types",
                "info": "Some sources could not be recognized automatically. Select a type for each row.",
                "headers": ["Original row", "Source type", "Issues"],
                "select_type_title": "Source type required",
                "select_type_text": "Select a source type for every row.",
                "types": {
                    "unknown": "Unknown", "book": "Book", "journal_article": "Journal article",
                    "collection_article": "Collection article", "multi_volume": "Multi-volume edition",
                    "gost_standard": "GOST / standard", "federal_law": "Federal law",
                    "dissertation": "Dissertation / abstract", "archive": "Archive document",
                    "electronic": "Electronic resource", "article": "Article (unspecified)",
                },
            },
            "zh_CN": {
                "title": "确认来源类型",
                "info": "部分来源无法自动识别。请为每一行选择类型。",
                "headers": ["原始行", "来源类型", "问题"],
                "select_type_title": "需要选择类型",
                "select_type_text": "请为所有行选择来源类型。",
                "types": {
                    "unknown": "未知", "book": "图书", "journal_article": "期刊文章",
                    "collection_article": "文集文章", "multi_volume": "多卷本",
                    "gost_standard": "GOST / 标准", "federal_law": "联邦法律",
                    "dissertation": "论文 / 摘要", "archive": "档案文档",
                    "electronic": "电子资源", "article": "文章（未细分）",
                },
            },
        }
        return texts.get(self.lang, texts["ru_RU"])[key]

    def _setup_ui(self):
        layout = QVBoxLayout(self)

        info = QLabel(self._t("info"))
        layout.addWidget(info)

        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(self._t("headers"))
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.table.verticalHeader().setVisible(False)
        self.table.verticalHeader().setDefaultSectionSize(34)
        self.table.setRowCount(len(self.issue_rows))
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        layout.addWidget(self.table)

        colors = self._theme_colors()
        bg = QColor(colors["input"])
        fg = QColor(colors["text"])
        for row, item in enumerate(self.issue_rows):
            idx = item["index"]
            text = item["text"]
            missing = item.get("missing", [])
            source_item = QTableWidgetItem(text)
            source_item.setBackground(bg)
            source_item.setForeground(fg)
            self.table.setItem(row, 0, source_item)
            combo = QComboBox()
            for code, label in self.TYPE_OPTIONS:
                combo.addItem(self._t("types").get(code, label), code)
            current_type = item.get("type", "unknown")
            combo.setCurrentIndex(max(0, combo.findData(current_type)))
            combo.setMinimumWidth(220)
            self.table.setCellWidget(row, 1, combo)
            missing_item = QTableWidgetItem(", ".join(missing))
            missing_item.setBackground(bg)
            missing_item.setForeground(fg)
            self.table.setItem(row, 2, missing_item)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self._on_accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def _on_accept(self):
        self._type_map = {}
        for row, item in enumerate(self.issue_rows):
            idx = item["index"]
            combo = self.table.cellWidget(row, 1)
            type_code = combo.currentData()
            if type_code == "unknown":
                QMessageBox.warning(self, self._t("select_type_title"), self._t("select_type_text"))
                return
            self._type_map[idx] = type_code

        self.accept()

    def get_type_map(self) -> dict:
        return self._type_map


class LogViewerDialog(QDialog):
    """Просмотр логов ИИ"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Логи ИИ")
        self.setMinimumSize(900, 600)
        self._setup_ui()
        self._load_logs()

    def _setup_ui(self):
        layout = QVBoxLayout(self)

        self.tabs = QTabWidget()
        self.raw_log = QTextEdit()
        self.raw_log.setReadOnly(True)
        self.err_log = QTextEdit()
        self.err_log.setReadOnly(True)
        self.tabs.addTab(self.raw_log, "Ответы ИИ")
        self.tabs.addTab(self.err_log, "Ошибки")
        layout.addWidget(self.tabs)

        btns = QHBoxLayout()
        self.btn_refresh = QPushButton("Обновить")
        self.btn_refresh.clicked.connect(self._load_logs)
        btns.addWidget(self.btn_refresh)
        btns.addStretch()
        layout.addLayout(btns)

    def _load_logs(self):
        log_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "data", "logs"))
        raw_path = os.path.join(log_dir, "ai_raw_responses.log")
        err_path = os.path.join(log_dir, "ai_errors.log")

        self.raw_log.setText(self._read_log_file(raw_path))
        self.err_log.setText(self._read_log_file(err_path))

    @staticmethod
    def _read_log_file(path: str) -> str:
        if not os.path.exists(path):
            return f"Файл не найден: {path}"
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            return f"Ошибка чтения: {e}"


class BiblioMakerWindow(QMainWindow):
    """Главное окно приложения"""

    def __init__(self):
        super().__init__()
        self.db = GOSTDatabase()
        self.settings = self.db.get_all_settings()
        self.current_file_path: Optional[str] = None
        self.processing_thread: Optional[ProcessingThread] = None
        self._last_result: Optional[dict] = None

        self.setup_ui()
        self.load_history()

    def _t(self, key: str) -> str:
        translations = {
            "ru_RU": {
                "title": "BiblioMaker — Оформление по ГОСТ",
                "process_tab": "Обработка",
                "history_tab": "История",
                "ready": "Готов к работе",
                "open_document": "Открыть документ",
                "provider": "ИИ-провайдер:",
                "process": "Исправить по ГОСТ",
                "cancel": "Отмена",
                "save": "Сохранить",
                "settings": "Настройки",
                "input_source": "Источник данных",
                "manual_input": "Ручной ввод",
                "docx_file": "Файл .docx",
                "original_list": "Исходный список",
                "fixed_list": "Исправленный список",
                "history_headers": ["Дата", "Провайдер", "Исходный текст", "Результат"],
                "refresh": "Обновить",
                "clear_history": "Очистить историю",
                "logs": "Логи",
                "manual_placeholder": "Введите список литературы, каждый источник с новой строки:\n\nИванов А.А. Название книги. М.: Изд-во, 2020.\nПетров Б.Б. Статья // Журнал. 2021. № 3. С. 10-15.\nСидоров В.В. Сайт [Электронный ресурс]. URL: https://example.com",
            },
            "en_US": {
                "title": "BiblioMaker — GOST Formatting",
                "process_tab": "Processing",
                "history_tab": "History",
                "ready": "Ready",
                "open_document": "Open document",
                "provider": "AI provider:",
                "process": "Fix by GOST",
                "cancel": "Cancel",
                "save": "Save",
                "settings": "Settings",
                "input_source": "Data source",
                "manual_input": "Manual input",
                "docx_file": ".docx file",
                "original_list": "Original list",
                "fixed_list": "Fixed list",
                "history_headers": ["Date", "Provider", "Original text", "Result"],
                "refresh": "Refresh",
                "clear_history": "Clear history",
                "logs": "Logs",
                "manual_placeholder": "Enter the bibliography, one source per line:\n\nIvanov A.A. Book title. Moscow: Publisher, 2020.\nPetrov B.B. Article // Journal. 2021. No. 3. P. 10-15.\nSidorov V.V. Site [Electronic resource]. URL: https://example.com",
            },
            "zh_CN": {
                "title": "BiblioMaker — GOST 格式化",
                "process_tab": "处理",
                "history_tab": "历史",
                "ready": "就绪",
                "open_document": "打开文档",
                "provider": "AI 提供商:",
                "process": "按 GOST 修正",
                "cancel": "取消",
                "save": "保存",
                "settings": "设置",
                "input_source": "数据源",
                "manual_input": "手动输入",
                "docx_file": ".docx 文件",
                "original_list": "原始列表",
                "fixed_list": "修正列表",
                "history_headers": ["日期", "提供商", "原始文本", "结果"],
                "refresh": "刷新",
                "clear_history": "清空历史",
                "logs": "日志",
                "manual_placeholder": "请输入参考文献列表，每行一个来源:\n\nIvanov A.A. Book title. Moscow: Publisher, 2020.\nPetrov B.B. Article // Journal. 2021. No. 3. P. 10-15.\nSidorov V.V. Site [Electronic resource]. URL: https://example.com",
            },
        }
        lang = self.settings.get("ui_language", "ru_RU")
        return translations.get(lang, translations["ru_RU"]).get(key, translations["ru_RU"].get(key, key))

    @staticmethod
    def _fit_button_text(button: QPushButton, extra_padding: int = 46):
        width = button.fontMetrics().horizontalAdvance(button.text()) + extra_padding
        button.setMinimumWidth(max(button.minimumWidth(), width))
        button.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed)

    def setup_ui(self):
        """Настройка интерфейса"""
        self.setWindowTitle(self._t("title"))
        self.setMinimumSize(1200, 800)
        self.setFont(QFont("Segoe UI Variable", int(self.settings.get("font_size", "12"))))
        self._apply_styles()

        # Центральный виджет
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)

        # Верхняя панель
        top_panel = self._create_top_panel()
        layout.addLayout(top_panel)

        # Вкладки
        self.tabs = QTabWidget()

        # Вкладка 1: Обработка
        tab_process = self._create_process_tab()
        self.tabs.addTab(tab_process, self._t("process_tab"))

        # Вкладка 2: История
        tab_history = self._create_history_tab()
        self.tabs.addTab(tab_history, self._t("history_tab"))

        layout.addWidget(self.tabs)

        # Панель прогресса
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        # Статус бар
        self.statusLabel = QLabel(self._t("ready"))
        self.statusBar().addWidget(self.statusLabel)
        self._set_input_mode('manual')
        self._apply_settings_to_widgets()

    def _create_top_panel(self):
        """Создание верхней панели"""
        panel = QHBoxLayout()
        panel.setContentsMargins(0, 0, 0, 0)
        panel.setSpacing(12)

        # Кнопка загрузки
        self.btn_load = QPushButton(self._t("open_document"))
        self.btn_load.clicked.connect(self.load_document)
        self.btn_load.setMinimumHeight(40)
        panel.addWidget(self.btn_load)

        # Выбор провайдера
        self.provider_label = QLabel(self._t("provider"))
        panel.addWidget(self.provider_label)
        self.provider_combo = QComboBox()
        self.provider_combo.addItems(['mock', 'yandex'])
        panel.addWidget(self.provider_combo)

        # Кнопка обработки
        btn_process = QPushButton(self._t("process"))
        btn_process.setObjectName("primary")
        btn_process.setMinimumHeight(40)
        btn_process.clicked.connect(self.process_document)
        btn_process.setEnabled(True)
        self.btn_process = btn_process
        panel.addWidget(btn_process)

        # Кнопка отмены
        btn_cancel = QPushButton(self._t("cancel"))
        btn_cancel.setObjectName("danger")
        btn_cancel.setMinimumHeight(40)
        btn_cancel.clicked.connect(self.cancel_processing)
        btn_cancel.setEnabled(False)
        self.btn_cancel = btn_cancel
        panel.addWidget(btn_cancel)

        # Кнопка сохранения
        btn_save = QPushButton(self._t("save"))
        btn_save.clicked.connect(self.save_document)
        btn_save.setEnabled(False)
        self.btn_save = btn_save
        panel.addWidget(btn_save)

        # Настройки
        self.btn_settings = QPushButton(self._t("settings"))
        self.btn_settings.clicked.connect(self.open_settings)
        panel.addWidget(self.btn_settings)

        panel.addStretch()
        return panel

    def _create_process_tab(self):
        """Создание вкладки обработки"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        # === Переключатель режима ввода ===
        self.input_mode_group = QGroupBox(self._t("input_source"))
        input_mode_layout = QHBoxLayout(self.input_mode_group)
        input_mode_layout.setContentsMargins(12, 16, 12, 12)
        input_mode_layout.setSpacing(8)
        self.input_mode_group.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.input_mode_group.setMaximumHeight(90)

        self.radio_manual = QPushButton(self._t("manual_input"))
        self.radio_manual.setCheckable(True)
        self.radio_manual.clicked.connect(lambda: self._set_input_mode('manual'))
        self.radio_manual.setMinimumWidth(120)

        self.radio_file = QPushButton(self._t("docx_file"))
        self.radio_file.setCheckable(True)
        self.radio_file.clicked.connect(lambda: self._set_input_mode('file'))
        self.radio_file.setMinimumWidth(110)
        self.radio_manual.setChecked(True)

        button_group = QWidget()
        btn_layout = QHBoxLayout(button_group)
        btn_layout.setContentsMargins(0, 0, 0, 0)
        btn_layout.setSpacing(8)
        btn_layout.addWidget(self.radio_manual)
        btn_layout.addWidget(self.radio_file)

        input_mode_layout.addWidget(button_group)
        input_mode_layout.addStretch()
        layout.addWidget(self.input_mode_group)
        # === Конец переключателя ===

        # Сплиттер
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setHandleWidth(1)

        # Левая панель: исходный
        self.left_group = QGroupBox(self._t("original_list"))
        left_layout = QVBoxLayout(self.left_group)
        left_layout.setContentsMargins(12, 16, 12, 12)
        left_layout.setSpacing(8)

        # Текстовое поле для ручного ввода
        self.manual_input = QTextEdit()
        self.manual_input.setPlaceholderText(self._t("manual_placeholder"))
        self.manual_input.setFont(QFont("Cascadia Mono", 10))
        self.manual_input.setVisible(False)
        self.manual_input.textChanged.connect(self._on_manual_text_changed)
        left_layout.addWidget(self.manual_input)

        # Поле для просмотра из файла
        self.source_text = QTextEdit()
        self.source_text.setReadOnly(True)
        self.source_text.setFont(QFont("Cascadia Mono", 10))
        left_layout.addWidget(self.source_text)

        splitter.addWidget(self.left_group)

        # Правая панель: исправленный
        self.right_group = QGroupBox(self._t("fixed_list"))
        right_layout = QVBoxLayout(self.right_group)
        right_layout.setContentsMargins(12, 16, 12, 12)
        right_layout.setSpacing(8)
        self.fixed_text = QTextEdit()
        self.fixed_text.setReadOnly(True)
        self.fixed_text.setFont(QFont("Cascadia Mono", 10))
        right_layout.addWidget(self.fixed_text)
        splitter.addWidget(self.right_group)

        splitter.setSizes([500, 500])
        layout.addWidget(splitter)

        return widget

    def _create_history_tab(self):
        """Создание вкладки истории"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        # Таблица истории
        self.history_table = QTableWidget()
        self.history_table.setColumnCount(4)
        self.history_table.setHorizontalHeaderLabels([
            *self._t("history_headers")
        ])
        header = self.history_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        self.history_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.history_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectItems)
        self.history_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.history_table.customContextMenuRequested.connect(self._show_history_context_menu)
        layout.addWidget(self.history_table)

        # Кнопки
        btn_layout = QHBoxLayout()

        self.btn_refresh_history = QPushButton(self._t("refresh"))
        self.btn_refresh_history.clicked.connect(self.load_history)
        btn_layout.addWidget(self.btn_refresh_history)

        self.btn_clear_history = QPushButton(self._t("clear_history"))
        self.btn_clear_history.clicked.connect(self.clear_history)
        btn_layout.addWidget(self.btn_clear_history)

        self.btn_logs = QPushButton(self._t("logs"))
        self.btn_logs.clicked.connect(self.open_logs)
        btn_layout.addWidget(self.btn_logs)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        return widget

    def _set_input_mode(self, mode: str):
        """Переключение между режимами ввода"""
        if mode == 'file':
            self.radio_file.setChecked(True)
            self.radio_manual.setChecked(False)
            self.manual_input.setVisible(False)
            self.source_text.setVisible(True)
        else:  # manual
            self.radio_file.setChecked(False)
            self.radio_manual.setChecked(True)
            self.manual_input.setVisible(True)
            self.source_text.setVisible(False)

    def _on_manual_text_changed(self):
        """Обработчик изменения текста в ручном режиме"""
        pass

    def _get_input_text(self) -> List[str]:
        """Получение текста из текущего источника"""
        if self.radio_file.isChecked() and self.current_file_path:
            processor = DocxProcessor(self.current_file_path)
            return processor.extract_bibliography()
        else:
            text = self.manual_input.toPlainText().strip()
            if not text:
                return []
            return [line.strip() for line in text.split('\n') if line.strip()]

    def load_document(self):
        """Загрузка документа"""
        self._set_input_mode('file')

        file_path, _ = QFileDialog.getOpenFileName(
            self, "Открыть документ", "", "Word Documents (*.docx)"
        )

        if file_path:
            self.current_file_path = file_path
            self.statusLabel.setText(f"Загружен: {os.path.basename(file_path)}")

            try:
                processor = DocxProcessor(file_path)
                entries = processor.extract_bibliography()
                self.source_text.setText('\n'.join(entries))
                self.btn_process.setEnabled(len(entries) > 0)
                self.statusLabel.setText(
                    f"Загружен: {os.path.basename(file_path)} | "
                    f"Найдено источников: {len(entries)}"
                )
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось прочитать документ: {e}")

    def process_document(self):
        """Обработка документа"""
        entries = self._get_input_text()

        if not entries:
            QMessageBox.warning(
                self, "Нет данных",
                "Введите список литературы или загрузите .docx, затем повторите."
            )
            return

        self.source_text.setText('\n'.join(entries))

        analysis = [analyze_entry(e) for e in entries]
        auto_types = [t for t, _ in analysis]
        issue_rows = []
        for i, (t, missing) in enumerate(analysis):
            if t == "unknown" or missing:
                issue_rows.append({
                    "index": i,
                    "text": entries[i],
                    "type": t,
                    "missing": missing
                })

        if issue_rows:
            dialog = TypeResolutionDialog(self, issue_rows)
            if dialog.exec() != QDialog.DialogCode.Accepted:
                return
            overrides = dialog.get_type_map()
        else:
            overrides = {}

        prompt_entries = []
        for i, line in enumerate(entries):
            t = overrides.get(i, auto_types[i])
            prompt_entries.append(f"TYPE={t}: {line}")

        provider_name = self.provider_combo.currentText()

        self.processing_thread = ProcessingThread(
            file_path=self.current_file_path if self.radio_file.isChecked() else None,
            entries=entries,
            prompt_entries=prompt_entries,
            provider_name=provider_name,
            ui_language=self.settings.get("ui_language", "ru_RU")
        )
        self.processing_thread.progress.connect(self.update_progress)
        self.processing_thread.finished.connect(self.processing_finished)
        self.processing_thread.error.connect(self.processing_error)
        self.processing_thread.cancelled.connect(self.processing_cancelled)

        self.progress_bar.setVisible(True)
        self.btn_process.setEnabled(False)
        self.btn_cancel.setEnabled(True)
        self.processing_thread.start()

    def update_progress(self, value: int, message: str):
        """Обновление прогресса"""
        self.progress_bar.setValue(value)
        self.statusLabel.setText(message)

    def processing_finished(self, result: dict):
        """Завершение обработки"""
        self.progress_bar.setVisible(False)
        self.btn_process.setEnabled(True)
        self.btn_cancel.setEnabled(False)
        self.btn_save.setEnabled(True)
        self._last_result = result

        self.fixed_text.setText('\n'.join(result['fixed']))

        # Сохранение результата в БД
        self.db.add_history_record(
            document_name="Ручной ввод" if result.get('is_manual') else os.path.basename(self.current_file_path or ""),
            document_path="" if result.get('is_manual') else (self.current_file_path or ""),
            ai_provider=result['provider'],
            sources_count=result['sources_count'],
            changes_count=result['changes_count'],
            original_text='\n'.join(result['original']),
            fixed_text='\n'.join(result['fixed']),
            status='Успешно',
            input_mode='manual' if result.get('is_manual') else 'file'
        )

        self.load_history()

        self.statusLabel.setText(
            f"Обработано! Изменений: {result['changes_count']} | "
            f"Провайдер: {result['provider']}"
        )

        QMessageBox.information(
            self, "Готово",
            f"Обработка завершена!\n\n"
            f"Источников: {result['sources_count']}\n"
            f"Изменений: {result['changes_count']}\n"
            f"Провайдер: {result['provider']}"
        )

    def processing_error(self, error: str):
        """Ошибка обработки"""
        self.progress_bar.setVisible(False)
        self.btn_process.setEnabled(True)
        self.btn_cancel.setEnabled(False)

        QMessageBox.critical(self, "Ошибка", error)
        self.statusLabel.setText("Ошибка обработки")

    def processing_cancelled(self, message: str):
        """Отмена обработки"""
        self.progress_bar.setVisible(False)
        self.btn_process.setEnabled(True)
        self.btn_cancel.setEnabled(False)

        QMessageBox.information(self, "Отмена", message)
        self.statusLabel.setText("Обработка отменена")

    def cancel_processing(self):
        """Запрос на отмену обработки"""
        if self.processing_thread and self.processing_thread.isRunning():
            self.statusLabel.setText("Отмена обработки...")
            self.btn_cancel.setEnabled(False)
            self.processing_thread.request_cancel()

    def save_document(self):
        """Сохранение результата"""
        fixed_text = self.fixed_text.toPlainText().strip()
        if not fixed_text:
            QMessageBox.warning(self, "Нет данных", "Нечего сохранять!")
            return

        output_folder = self.settings.get("output_folder", "").strip()
        initial_dir = output_folder if output_folder and os.path.isdir(output_folder) else ""
        if self.current_file_path and self.radio_file.isChecked():
            base_name = os.path.splitext(os.path.basename(self.current_file_path))[0]
            initial_dir = os.path.join(initial_dir or os.path.dirname(self.current_file_path), f"{base_name}_bibliomaker.docx")

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить результат", initial_dir,
            "Текстовые файлы (*.txt);;Word Documents (*.docx);;All Files (*)"
        )

        if not file_path:
            return

        try:
            if file_path.endswith('.docx'):
                lines = [line.strip() for line in fixed_text.split('\n') if line.strip()]
                if self.current_file_path and self.radio_file.isChecked():
                    if self.settings.get("auto_backup", "true").lower() == "true":
                        DocxProcessor.create_backup(self.current_file_path)
                    processor = DocxProcessor(self.current_file_path)
                    processor.replace_bibliography(lines, file_path)
                else:
                    from docx import Document
                    from docx.shared import Cm, Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    doc = Document()
                    heading = doc.add_heading("Список литературы (по ГОСТ)", level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for line in lines:
                        para = doc.add_paragraph(line)
                        para.paragraph_format.first_line_indent = Cm(1.25)
                        para.paragraph_format.line_spacing = Pt(14)

                    doc.save(file_path)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(fixed_text)

            QMessageBox.information(self, "Успех", "Результат сохранён!")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить: {e}")

    def load_history(self):
        """Загрузка истории из БД"""
        history = self.db.get_history(limit=50)

        self.history_table.setRowCount(len(history))

        for i, record in enumerate(history):
            self.history_table.setItem(i, 0, QTableWidgetItem(record['processed_at']))
            self.history_table.setItem(i, 1, QTableWidgetItem(record['ai_provider']))
            self.history_table.setItem(i, 2, QTableWidgetItem(record.get('original_text', '')))
            self.history_table.setItem(i, 3, QTableWidgetItem(record.get('fixed_text', '')))

    def clear_history(self):
        """Очистка истории"""
        reply = QMessageBox.question(
            self, "Подтверждение",
            "Вы уверены, что хотите очистить всю историю?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.db.clear_history()
            self.load_history()
            self.statusLabel.setText("История очищена")

    def open_settings(self):
        """Открытие диалога настроек"""
        dialog = SettingsDialog(self, self.db)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.settings = self.db.get_all_settings()
            self._apply_settings_to_widgets()

    def _apply_settings_to_widgets(self):
        """Применение сохраненных настроек к открытому окну."""
        self.setWindowTitle(self._t("title"))
        font_size = int(self.settings.get("font_size", "12"))
        self.setFont(QFont("Segoe UI Variable", font_size))
        self._apply_styles()

        provider = self.settings.get("default_provider", "mock")
        provider_index = self.provider_combo.findText(provider)
        if provider_index >= 0:
            self.provider_combo.setCurrentIndex(provider_index)

        self.tabs.setTabText(0, self._t("process_tab"))
        self.tabs.setTabText(1, self._t("history_tab"))
        self.btn_load.setText(self._t("open_document"))
        self.provider_label.setText(self._t("provider"))
        self.btn_process.setText(self._t("process"))
        self.btn_cancel.setText(self._t("cancel"))
        self.btn_save.setText(self._t("save"))
        self.btn_settings.setText(self._t("settings"))
        self.input_mode_group.setTitle(self._t("input_source"))
        self.radio_manual.setText(self._t("manual_input"))
        self.radio_file.setText(self._t("docx_file"))
        self.left_group.setTitle(self._t("original_list"))
        self.right_group.setTitle(self._t("fixed_list"))
        self.manual_input.setPlaceholderText(self._t("manual_placeholder"))
        self.history_table.setHorizontalHeaderLabels(self._t("history_headers"))
        self.btn_refresh_history.setText(self._t("refresh"))
        self.btn_clear_history.setText(self._t("clear_history"))
        self.btn_logs.setText(self._t("logs"))
        for button in (self.radio_manual, self.radio_file):
            self._fit_button_text(button)
        if self.statusLabel.text() in ("Готов к работе", "Ready", "就绪"):
            self.statusLabel.setText(self._t("ready"))

    def open_logs(self):
        """Открытие диалога логов"""
        dialog = LogViewerDialog(self)
        dialog.exec()

    def _show_history_context_menu(self, position):
        """Контекстное меню для копирования"""
        item = self.history_table.itemAt(position)
        if not item:
            return

        menu = QMenu(self)
        action_copy = QAction("Копировать", self)
        action_copy.triggered.connect(lambda: self._copy_history_cell(item))
        menu.addAction(action_copy)
        menu.exec(self.history_table.viewport().mapToGlobal(position))

    def _copy_history_cell(self, item: QTableWidgetItem):
        """Копирование текста ячейки"""
        clipboard = QApplication.clipboard()
        clipboard.setText(item.text())

    def _apply_styles(self):
        """Единый современный стиль интерфейса"""
        theme = self.settings.get("theme", "light")
        font_size = int(self.settings.get("font_size", "12"))
        if theme == "dark":
            colors = {
                "window": "#111827",
                "text": "#e5e7eb",
                "disabled_text": "#6b7280",
                "pane": "#1f2937",
                "muted": "#9ca3af",
                "tab": "#374151",
                "border": "#4b5563",
                "button": "#1f2937",
                "button_hover": "#374151",
                "button_pressed": "#4b5563",
                "input": "#111827",
                "header": "#374151",
                "progress": "#374151",
                "selection": "#1d4ed8",
                "selection_text": "#ffffff",
                "grid": "#4b5563",
                "scroll": "#6b7280",
            }
        else:
            colors = {
                "window": "#f6f7fb",
                "text": "#0f172a",
                "disabled_text": "#94a3b8",
                "pane": "#ffffff",
                "muted": "#475569",
                "tab": "#e9edf5",
                "border": "#d7dde7",
                "button": "#ffffff",
                "button_hover": "#f1f5f9",
                "button_pressed": "#e2e8f0",
                "input": "#ffffff",
                "header": "#f1f5f9",
                "progress": "#e2e8f0",
                "selection": "#bfdbfe",
                "selection_text": "#0f172a",
                "grid": "#e2e8f0",
                "scroll": "#cbd5e1",
            }

        stylesheet = """
            QMainWindow, QDialog, QMessageBox {{
                background: {window};
            }}
            QWidget, QLabel, QCheckBox {{
                color: {text};
                font-size: {font_size}pt;
            }}
            QWidget:disabled, QLabel:disabled, QCheckBox:disabled {{
                color: {disabled_text};
            }}
            QTabWidget::pane {{
                border: 1px solid {border};
                border-top: none;
                background: {pane};
                border-radius: 10px;
            }}
            QTabBar::tab {{
                background: {tab};
                border: 1px solid {border};
                border-bottom: none;
                padding: 10px 16px;
                margin-right: 6px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }}
            QTabBar::tab:selected {{
                background: {pane};
                border-color: {border};
            }}
            QGroupBox {{
                border: 1px solid {border};
                border-radius: 10px;
                margin-top: 12px;
                background: {pane};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 0 6px;
                margin-left: 10px;
                color: {muted};
                background: {pane};
            }}
            QPushButton {{
                background: {button};
                border: 1px solid {border};
                border-radius: 8px;
                padding: 8px 14px;
            }}
            QPushButton:hover {{
                background: {button_hover};
            }}
            QPushButton:pressed {{
                background: {button_pressed};
            }}
            QPushButton:checked {{
                background: {button_pressed};
                border: 1px solid {border};
                font-weight: 600;
            }}
            QPushButton:disabled {{
                color: {disabled_text};
                background: {tab};
            }}
            QPushButton#primary {{
                background: #2563eb;
                color: #ffffff;
                border: 1px solid #1d4ed8;
                font-weight: 600;
            }}
            QPushButton#danger {{
                background: #ef4444;
                color: #ffffff;
                border: 1px solid #dc2626;
                font-weight: 600;
            }}
            QLineEdit, QTextEdit, QPlainTextEdit, QComboBox, QSpinBox {{
                background: {input};
                border: 1px solid {border};
                border-radius: 8px;
                padding: 6px 8px;
                selection-background-color: {selection};
                selection-color: {selection_text};
                color: {text};
            }}
            QLineEdit:disabled, QTextEdit:disabled, QPlainTextEdit:disabled, QComboBox:disabled, QSpinBox:disabled {{
                color: {disabled_text};
                background: {pane};
            }}
            QComboBox::drop-down {{
                border: none;
                width: 24px;
            }}
            QComboBox QAbstractItemView {{
                background: {pane};
                color: {text};
                border: 1px solid {border};
                selection-background-color: {selection};
                selection-color: {selection_text};
                outline: none;
            }}
            QSpinBox::up-button, QSpinBox::down-button {{
                background: {button};
                border: 1px solid {border};
                width: 18px;
            }}
            QSpinBox::up-button:hover, QSpinBox::down-button:hover {{
                background: {button_hover};
            }}
            QMenu {{
                background: {pane};
                color: {text};
                border: 1px solid {border};
                padding: 4px;
            }}
            QMenu::item {{
                padding: 6px 24px 6px 12px;
                border-radius: 4px;
            }}
            QMenu::item:selected {{
                background: {selection};
                color: {selection_text};
            }}
            QTableWidget, QTableView {{
                background: {input};
                alternate-background-color: {pane};
                color: {text};
                border: 1px solid {border};
                gridline-color: {grid};
                selection-background-color: {selection};
                selection-color: {selection_text};
            }}
            QTableWidget::item, QTableView::item {{
                background: {input};
                color: {text};
                border-color: {grid};
                padding: 4px;
            }}
            QTableWidget::item:selected, QTableView::item:selected {{
                background: {selection};
                color: {selection_text};
            }}
            QHeaderView::section {{
                background: {header};
                border: 1px solid {border};
                color: {text};
                padding: 6px;
            }}
            QTableCornerButton::section {{
                background: {header};
                border: 1px solid {border};
            }}
            QStatusBar {{
                background: {window};
                color: {text};
            }}
            QProgressBar {{
                background: {progress};
                border: none;
                border-radius: 6px;
                height: 12px;
                text-align: center;
                color: {text};
            }}
            QProgressBar::chunk {{
                background: #2563eb;
                border-radius: 6px;
            }}
            QSplitter::handle:horizontal {{
                background: {border};
            }}
            QScrollBar:vertical, QScrollBar:horizontal {{
                background: {pane};
                border: none;
                width: 14px;
                height: 14px;
                margin: 0;
            }}
            QScrollBar::handle:vertical, QScrollBar::handle:horizontal {{
                background: {scroll};
                border-radius: 6px;
                min-height: 24px;
                min-width: 24px;
            }}
            QScrollBar::handle:vertical:hover, QScrollBar::handle:horizontal:hover {{
                background: {muted};
            }}
            QScrollBar::add-line, QScrollBar::sub-line {{
                width: 0;
                height: 0;
            }}
            QScrollBar::add-page, QScrollBar::sub-page {{
                background: transparent;
            }}
        """.format(font_size=font_size, **colors)
        app = QApplication.instance()
        if app is not None:
            app.setStyleSheet(stylesheet)
        self.setStyleSheet(stylesheet)
